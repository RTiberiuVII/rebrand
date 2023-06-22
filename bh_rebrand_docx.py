'''
Author: Ole Gerlof
Amended by: Tiberiu Rociu
Date: 21/06/2023
Version: 1.3
Purpose:
Rebrand MS Office documents (Word, Excel PowerPoint) from
    former â€œBHGEâ€ to current Baker Hughes Company style and re-create PDF files
Applies to company logo, company name, document font, brand colors

'''

import sys
import os
import re
from time import time
from datetime import datetime
from re import sub, search, findall, IGNORECASE
from zipfile import ZipFile, ZIP_DEFLATED
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import comtypes.client
from docx.table import _Cell
import xml.etree.ElementTree as ET

FILE_FORMAT_PDF_WORD = 17
FILE_FORMAT_PDF_EXCEL = 0
FILE_FORMAT_PDF_PPT = 2


def count_docx(file_name):
    document = Document(file_name)

    newparatextlist = []
    for paratext in document.paragraphs:
        newparatextlist.append(paratext.text)

    return len(re.findall(r'\w+', '\n'.join(newparatextlist)))


def shuttle_text(shuttle):
    t = ''
    for i in shuttle:
        t += i.text
    return t


def replace_text(runs, target, replace):
    """
    Replaces the target string with the replace string in the runs.

    Parameters:
    - runs (List[Run]): List of runs to search and replace within.
    - target (str): The string to be replaced.
    - replace (str): The replacement string.
    """
    begin = 0
    full_text = ""
    for end, run in enumerate(runs):
        full_text += run.text
        # print('Full text: ', full_text)
        if target in full_text:
            # print('Target in full text. Target: ', target, ' full text: ', full_text)
            # Find the beginning index
            index = full_text.index(target)
            while index >= len(runs[begin].text):
                index -= len(runs[begin].text)
                begin += 1

            # Find the corresponding runs
            shuttle = runs[begin:end + 1]

            # Perform the replace operation
            if target in shuttle[0].text:
                # print('Replacing: ', target, ' with: ', shuttle[0].text)
                shuttle[0].text = shuttle[0].text.replace(target, replace)
            else:
                replace_begin_index = full_text.index(target)
                replace_end_index = replace_begin_index + len(target)
                replace_end_index_in_last_run = replace_end_index - len(''.join(run.text for run in shuttle[:-1]))
                shuttle[0].text = shuttle[0].text[:replace_begin_index] + replace

                # Clear middle runs
                for i in shuttle[1:-1]:
                    i.text = ''

                # Keep last run
                shuttle[-1].text = shuttle[-1].text[replace_end_index_in_last_run:]

            # Reset the begin index for the next iteration
            begin = end + 1
            full_text = ""
    return


def docx_replace(doc, target, replace):

    # Replace text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Replace text in table cells
                    replace_text(paragraph.runs, target, replace)
                    # Replace text in hyperlinks within table cells
                    for link in paragraph._element.xpath(".//w:hyperlink"):
                        replace_text(link.xpath("w:r", namespaces=link.nsmap), target, replace)

    # Replace text in headers and footers
    for section in doc.sections:
        # Process the header and footer for the main body and first page
        for header in [section.header, section.first_page_header]:
            for paragraph in header.paragraphs:
                # Find all <w:t> tags within the header paragraphs
                w_t_tags = paragraph._element.xpath(".//w:t")

                # Replace text in header paragraphs
                replace_text(w_t_tags, target, replace)

            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Replace text in table cells within headers
                            replace_text(paragraph.runs, target, replace)
                            # Replace text in hyperlinks within table cells within headers
                            for link in paragraph._element.xpath(".//w:hyperlink"):
                                replace_text(link.xpath("w:r", namespaces=link.nsmap), target, replace)

        for footer in [section.footer, section.first_page_footer]:
            for paragraph in footer.paragraphs:
                # Replace text in footer paragraphs
                replace_text(paragraph.runs, target, replace)

            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            # Replace text in table cells within footers
                            replace_text(paragraph.runs, target, replace)
                            # Replace text in hyperlinks within table cells within footers
                            for link in paragraph._element.xpath(".//w:hyperlink"):
                                replace_text(link.xpath("w:r", namespaces=link.nsmap), target, replace)

    # Replace text in the main document body
    for paragraph in doc.paragraphs:
        # Replace text in main body paragraphs
        replace_text(paragraph.runs, target, replace)
        # Replace text in hyperlinks within main body paragraphs
        for link in paragraph._element.xpath(".//w:hyperlink"):
            replace_text(link.xpath("w:r", namespaces=link.nsmap), target, replace)


def text_rebrand(file_in, config):
    # Store file path from CL Arguments.
    file_path = file_in

    if file_path.endswith('.docx'):
        doc = Document(file_path)

        # initialize variable for tables
        tables = doc.tables
        # initialize variable for counting text match occurrences
        occurrences = {}

        for replace_duo in config["ReplaceString"]:
            # initialize the number of occurrences to 0
            occurrences[replace_duo[0]] = 0

            # Loop through sections for header footer content
            for section in doc.sections:

                # Check text of first page header
                fheader = section.first_page_header
                # fheader.is_linked_to_previous = True

                for fheaderpar in fheader.paragraphs:

                    # check if paragraph contains text
                    if fheaderpar.text:

                        # check if target text exists in paragraph text
                        if replace_duo[0] in fheaderpar.text:
                            text = fheaderpar.text.replace(replace_duo[0], replace_duo[1])

                            if text != fheaderpar.text:
                                # Replace the text and increment the number of occurrences
                                fheaderpar.text = text
                                occurrences[replace_duo[0]] += 1

                # Check text of first page footer
                ffooter = section.first_page_footer
                # ffooter.is_linked_to_previous = True

                for ffooterpar in ffooter.paragraphs:

                    # check if paragraph contains text
                    if ffooterpar.text:
                        'print(ffooterpar.text)'
                        # check if target text exists in paragraph text
                        if replace_duo[0] in ffooterpar.text:
                            text = ffooterpar.text.replace(replace_duo[0], replace_duo[1])

                            if text != ffooterpar.text:
                                # Replace the text and increment the number of occurrences
                                ffooterpar.text = text
                                occurrences[replace_duo[0]] += 1

        # Loop through tables in document
        for table in tables:

            # Loop through rows in table
            for row in table.rows:

                # Loop through cells in row
                for cell in row.cells:

                    # Loop through paragraphs in cell
                    for paragraph in cell.paragraphs:

                        # check if paragraph contains text
                        if paragraph.text:

                            # check if target text exists in paragraph text
                            if replace_duo[0] in paragraph.text:
                                text = paragraph.text.replace(replace_duo[0], replace_duo[1])

                                # Check if replaced text is not the same as original
                                if text != paragraph.text:

                                    # Replace the text and increment the number of occurrences
                                    paragraph.text = text
                                    occurrences[replace_duo[0]] += 1

        # Loop through paragraphs
        for para in doc.paragraphs:

            # check if paragraph contains text
            if para.text:

                # check if target text exists in paragraph text
                if replace_duo[0] in para.text:
                    # replaced_text = re.sub(str(replace_duo[0]), str(replace_duo[1]), run.text, 999)
                    # style = run.style
                    # run.style = style
                    text = para.text.replace(replace_duo[0], replace_duo[1])

                    # Check if replaced text is not the same as original
                    if text != para.text:

                        # Replace the text and increment the number of occurrences
                        para.text = text
                        occurrences[replace_duo[0]] += 1

            # Loop through hyperlinks in document
            for link in para._element.xpath(".//w:hyperlink"):
                inner_run = link.xpath("w:r", namespaces=link.nsmap)[0]

                # Check if hyperlink contains text
                if inner_run.text:

                    # Check if hyperlink text contains target tex
                    if replace_duo[0] in inner_run.text:
                        text = inner_run.text.replace(replace_duo[0], replace_duo[1])
                        'print(text)'

                        # Check if replaced text is not the same as original
                        if text != inner_run.text:

                            # Replace the text and increment the number of occurrences
                            inner_run.text = text
                            occurrences[replace_duo[0]] += 1

    # print the number of occurrences of each word
    for word, count in occurrences.items():
        if count > 0:
            print(f"{word} ({count})")

    # make a new file name by adding "_new" to the original file name
    # new_file_path = file_path.replace(".docx", "_new.docx")
    # save the new docx file
    new_file_path = os.path.basename(file_path)
    doc.save((config["OutputFolder"]) + new_file_path)

    return


def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph


def copy_and_replace(zip_in, zip_out, config):
    '''
    Copies the zip file except for the media folder

    Parameters
    ----------
    zip_in: ZipFile obj
        input ZipFile

    zip_out: ZipFile obj
        output ZipFile

    config: dictionary
        dict contain information from the config file
    '''
    # Go over every file in the input document
    textnote = ""
    for path in zip_in.namelist():
        # Check if the path is not in the media folder
        if not "media" in path:
            file_content = zip_in.read(path)
            decoded_content = try_decode(file_content)
            # Replace specified content
            for replace_duo in config["ReplaceString"]:
                # test = file_content.find(replace_duo[0])
                    # findall(replace_duo[0], file_content)
                # print(test)
                #if findall(replace_duo[0], file_content):
                # if replace_duo[0] in file_content:
                if search(replace_duo[0], decoded_content):
                    textnote += f"Legacy TextObject Found: {replace_duo[0]} in {path}, "
                '''
                if replace_duo[1] in config["case-unsensitive"]:
                    file_content = sub(replace_duo[0], replace_duo[1], file_content, flags=IGNORECASE)
                else:
                    file_content = sub(replace_duo[0], replace_duo[1], file_content)
                '''
            # Copy file into new document
            zip_out.writestr(path, file_content)
    return textnote

def try_decode(content):
    encodings = ['utf-8', 'latin-1']  # Add more encodings if necessary
    for encoding in encodings:
        try:
            decoded_content = content.decode(encoding)
            return decoded_content
        except UnicodeDecodeError:
            continue
    return None

def replace_header_images(zip_in, zip_out, config, note):
    '''
    Search for images in the header and replace them with the logo

    Parameters
    ----------
    zip_in: ZipFile obj
        input ZipFile

    zip_out: ZipFile obj
        output ZipFile

    config: dictionary
        dict containg information from the config file

    Returns
    -------
    note: string
        notes about the location of the logo

    warning: string
        warning if multiple header images have been replaced
    '''
    print('CALLED header images')
    warning = ""
    # Check every filename for header --> can contain information about images in the header
    # Possible filenames: header2.xml.rels, header3.xml.rels
    for file in zip_in.namelist():
        if "header" in file and "rel" in file:
            # pylint: disable=W1401
            # Disable error from using \S\s in a binary string which is needed for regex
            # Get every target image
            image_locations = findall(b'Target="media[\S\s]*?"', zip_in.read(file))
            print('image_locations', image_locations)
            if len(image_locations) > 1:
                print(">>Changing multiple header images!")
                warning = "Warning: Multiple images found in header"
            for image_location in image_locations:
                image_location = image_location[image_location.find(b"media"):-1].decode("ascii")
                # Get the full path: Path in target string is just media/[image]
                # Path is either word/media/[image] or xls/media/[image]
                try:
                    header_image_path = f"{config['filetype']}/{image_location}"
                    if not header_image_path in zip_out.namelist():
                        print(f"Replaced header image at {header_image_path}")
                        zip_out.write(config["NewLogoPath"], header_image_path)
                        note += f"Replaced header image {header_image_path},"
                except KeyError:
                    # Filetype is not in config
                    # Get path to the image by looping over the entire content
                    for location in zip_in.namelist():
                        if image_location in location:
                            if not image_location in zip_out.namelist():
                                print(f"Replaced header image at {location}")
                                zip_out.write(config["NewLogoPath"], location)
                                note += f"Replaced header image {location},"
                            break

    return note, warning


def check_header_images(zip_in, config, note):
    '''
    checks header images and adds information to note

    Parameters
    ----------
    zip_in: ZipFile obj
        input ZipFile

    config: dictionary
        dict containg information from the config file

    note: String
        Information about the logo location

    Returns
    -------
    note: String
        Information about the logo location and possbile alternatives
    '''
    alternative_old_logo_found = False
    # Found image at expected path but with different size
    # Check if image is in header
    for file in zip_in.namelist():
        if "header" in file and "rel" in file:
            # pylint: disable=W1401
            # Disable error from using \S\s in a binary string which is needed for regex
            # Get every target image
            image_locations = findall(b'Target="media[\S\s]*?"', zip_in.read(file))
            for image_location in image_locations:
                image_location = image_location[image_location.find(
                    b"media"):-1].decode("ascii")
                # Format image_location to the full path
                try:
                    image_location = f'{config["filetype"]}/{image_location}'
                except KeyError:
                    # filetype not in config; cant automatically create full image path
                    # Check every location to get the full path
                    for path in zip_in.namelist():
                        if image_location in path:
                            image_location = path
                            break

                alternative_old_logo_size = zip_in.getinfo(
                    image_location).file_size
                note += f'''Found alternative image in {image_location
                    } with size {alternative_old_logo_size} bytes, '''
                alternative_old_logo_found = True
    if not alternative_old_logo_found:
        note += "No image in header,"
    return note


def place_logo(zip_in, zip_out, config):
    '''
    Places the new BH logo in the word document

    Parameters
    ----------
    zip_in: ZipFile obj
        input ZipFile

    zip_out: ZipFile obj
        output ZipFile

    config: dictionary
        dict containg information from the config file

    Returns
    -------
    status: string
        status information about the finding of the logo

    note: string
        notes about the location of the logo

    warning: string
        warning if multiple header images have been replaced
    '''
    namelist = zip_in.namelist()
    status = "LogoNotFound"
    note = ""
    warning = ""
    headercount = 0
    imagecount = 0

    # Check expected logo path and logo size
    '''
    if config["OldLogoPath_formatted"] in namelist:
        if zip_in.getinfo(config["OldLogoPath_formatted"]).file_size == int(config["OldLogoSize"]):
            # Found image at expected path with expected size -> copy new logo
            print("Logo found at expected location")
            'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
            status = "OK"

    else:
        print("Logo not in expected location")
    
    '''
    for file in zip_in.namelist():
        if "header" in file and "rel" in file:
            # pylint: disable=W1401
            # Disable error from using \S\s in a binary string which is needed for regex
            # Get every target image
            headercount += 1
            image_locations = findall(b'Target="media[\S\s]*?"', zip_in.read(file))
            for image_location in image_locations:
                imagecount += 1
                image_location = image_location[image_location.find(
                    b"media"):-1].decode("ascii")
                # Format image_location to the full path
                try:
                    image_location = f'{config["filetype"]}/{image_location}'
                except KeyError:
                    # filetype not in config; cant automatically create full image path
                    # Check every location to get the full path
                    for path in zip_in.namelist():
                        if image_location in path:
                            image_location = path
                            break

    if headercount == 0:
        status = "No Header"

    if headercount > 0:
        status = "Header Found, "

    if imagecount == 1:
        status += "Single Header Image, "
        if image_location in namelist:
            if zip_in.getinfo(image_location).file_size == int(7093) \
                    or zip_in.getinfo(image_location).file_size == int(8362) \
                    or zip_in.getinfo(image_location).file_size == int(1395) \
                    or zip_in.getinfo(image_location).file_size == int(5184):
                # Found image at expected path with expected size -> copy new logo
                # print("Logo found at expected location")
                'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                status += "Legacy BH Logo"
            elif zip_in.getinfo(image_location).file_size == int(config["OldLogoSize"]) \
                    or zip_in.getinfo(image_location).file_size == int(29004):
                # Found image at expected path with expected size -> copy new logo
                # print("Logo found at expected location")
                'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                status += "GE Logo"
            elif zip_in.getinfo(image_location).file_size == int(35408) \
                    or zip_in.getinfo(image_location).file_size == int(8962) \
                            or zip_in.getinfo(image_location).file_size == int(119089) \
                            or zip_in.getinfo(image_location).file_size == int(65326) \
                            or zip_in.getinfo(image_location).file_size == int(47759):
                # Found image at expected path with new logo size
                'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                status += "Rebrand Logo"
            else:
                # Found image at expected path with new logo size
                zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])
                status += "Unknown Logo"

    if imagecount > 1:
        status += str(imagecount) + " Header Images, "
        multicount = 0
        for file in zip_in.namelist():
            if "header" in file and "rel" in file:
                image_locations = findall(b'Target="media[\S\s]*?"', zip_in.read(file))
                for image_location in image_locations:
                    multicount += 1
                    # status += str(multicount) + ", "
                    multi_image_location = image_location[image_location.find(
                        b"media"):-1].decode("ascii")
                    multi_image_location = f'{config["filetype"]}/{multi_image_location}'
                    if zip_in.getinfo(multi_image_location).file_size == int(7093) \
                            or zip_in.getinfo(multi_image_location).file_size == int(8362) \
                            or zip_in.getinfo(multi_image_location).file_size == int(1395) \
                            or zip_in.getinfo(multi_image_location).file_size == int(5184):
                        # Found image at expected path with expected size -> copy new logo
                        # print("Logo found at expected location")
                        'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                        status += "(" + str(multicount) + ") Legacy BH Logo, "
                    elif zip_in.getinfo(multi_image_location).file_size == int(config["OldLogoSize"]) \
                            or zip_in.getinfo(multi_image_location).file_size == int(29004):
                        # Found image at expected path with expected size -> copy new logo
                        # print("Logo found at expected location")
                        'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                        status += "(" + str(multicount) + ") GE Logo, "
                    elif zip_in.getinfo(multi_image_location).file_size == int(35408) \
                            or zip_in.getinfo(multi_image_location).file_size == int(8962) \
                            or zip_in.getinfo(multi_image_location).file_size == int(119089) \
                            or zip_in.getinfo(multi_image_location).file_size == int(65326) \
                            or zip_in.getinfo(multi_image_location).file_size == int(47759):
                        # Found image at expected path with new logo size
                        'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                        status += "(" + str(multicount) + ") Rebrand Logo, "
                    else:
                        # Found image at expected path with new logo size
                        'zip_out.write(config["NewLogoPath"], config["OldLogoPath_formatted"])'
                        status += "(" + str(multicount) + ") Unknown Logo, "

    print('config["ReplaceHeaderImage"].lower()', config["ReplaceHeaderImage"].lower())
    # Check if every image in the header should be replaced
    if "true" in config["ReplaceHeaderImage"].lower():
        print('Calling header images')
        note, warning = replace_header_images(zip_in, zip_out, config, note)
    else:
        note = check_header_images(zip_in, config, note)
    # Copy the rest of the images

    for path in namelist:
        if "media" in path:
            # Check for potential logo
            if (status == "LogoNotFound" and
                zip_in.getinfo(path).file_size == int(config["OldLogoSize"])):
                print(f"Replacing potential logo found at {path}")
                status = "LogoFoundInOtherPath"
                note += f'Found alternative logo in {path} with {config["OldLogoSize"]},'
            '''            
            if (status == "LogoNotFound" and
                zip_in.getinfo(path).file_size == int(config["OldLogoSize"])):
                if path not in zip_out.namelist():
                    # Potential logo found and logo was not at expected path -> copy new logo
                    print(f"Replacing potential logo found at {path}")
                    'zip_out.write(config["NewLogoPath"], path)'
                    status = "LogoFoundInOtherPath"
                    note += f'Found alternative logo in {path} with {config["OldLogoSize"]},'
            else:
                # Copy other images if they havent been copied
                if not path in zip_out.namelist():
                    'zip_out.writestr(path, zip_in.read(path))'
            '''
    return status, note, warning


def get_filetype(file, config):
    '''
    Returns the config containing information about the current filetype.
    When its neither a word file nor an excel file the config returned
    wont contain information about the filetype

    Parameters
    ----------
    file: string
        name of the file to be checked

    config: dictionary
        content of the configuration file as dictonary with the tag as
        key and the information as value

    Returns
    -------
    config: dictonary
        content of the configuration file as dictonary with the tag as
        key and the information as value with added filetype information
    '''
    if ".do" in file:
        config["filetype"] = "word"
    elif ".xls" in file:
        config["filetype"] = "xl"
    elif ".ppt" in file:
        config["filetype"] = "ppt"
    else:
        # Remove "filetype" from config if filetype is unknown
        try:
            config.pop("filetype")
        except KeyError:
            pass
    return config


def get_filetypes_in_folder(folder):
    '''
    Returns a list of all filetypes (Word, Excel or PowerPoint) in a folder

    Parameters
    ----------
    folder: string
        the path to the folder containing the files

    Returns
    -------
    filetype: list
        list of the filetypes (Word,Excel or PowerPoint) found in the folder
    '''
    filetypes = []
    for file in os.listdir(folder):
        if ".do" in file:
            if "Word" not in filetypes:
                filetypes.append("Word")
        elif ".xls" in file:
            if "Excel" not in filetypes:
                filetypes.append("Excel")
        elif ".ppt" in file:
            if "PowerPoint" not in filetypes:
                filetypes.append("PowerPoint")
    return filetypes


def start_com_servers(filetypes):
    '''
    Starts the needed COM servers for PDF conversion

    Parameters
    ----------
    filetype: list
        list of filetypes of the files which are going to be converted
    '''
    word = None
    excel = None
    power_point = None
    if "Word" in filetypes:
        print("Starting Word COM server")
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
    if "Excel" in filetypes:
        print("Starting Excel COM server")
        excel = comtypes.client.CreateObject("Excel.Application")
        excel.Visible = False
    if "PowerPoint" in filetypes:
        print("Starting PowerPoint COM server")
        power_point = comtypes.client.CreateObject("PowerPoint.Application")
        # Trying to hide the window will result in the program crashing
        #power_point.Visible = False
    return word, excel, power_point


def quit_com_servers(word, excel, power_point):
    '''
    Stops all active COM servers

    Parameters
    ----------
    word: comtypes.Pointer
        pointer to the Word COM server

    excel: comtypes.Pointer
        pointer to the Excel COM server

    power_point: comtypes.Pointer
        pointer to the PowerPoint COM server
    '''
    print("Stopping COM server")
    if word is not None:
        word.Quit()
    if excel is not None:
        excel.Quit()
    if power_point is not None:
        power_point.Quit()


def convert_to_pdf(file_in, config, word, excel, power_point):
    '''
    converts a file into a pdf file and saves it

    Parameters
    ----------
    file_in: string
        path to the input file

    config: dict
        contains information set in the config file and the current filetype

    word: comtypes.Pointer
        pointer to the Word COM server

    excel: comtypes.Pointer
        pointer to the Excel COM server

    power_point: comtypes.Pointer
        pointer to the PowerPoint COM server
    '''
    # Use the correct COM server
    if config["filetype"] == "word":
        doc = word.Documents.Open(file_in)
        doc.SaveAs(config["PDFPath"], FileFormat = FILE_FORMAT_PDF_WORD)
        doc.Close()
    elif config["filetype"] == "xl":
        workbook = excel.Workbooks.Open(file_in)
        workbook.ExportAsFixedFormat(FILE_FORMAT_PDF_EXCEL, config["PDFPath"])
        workbook.close()
    elif config["filetype"] == "ppt":
        presentation = power_point.Presentations.Open(file_in)
        presentation.ExportAsFixedFormat(config["PDFPath"], FILE_FORMAT_PDF_PPT)
        presentation.close()


def process_file(file_in, file_out, config, log):
    '''
    Replaces the old BHGE logo and copyright information

    Parameters
    ----------
    file_in: string
        the path to the input file

    file_out: string
        the path to the output file

    config: dictionary
        content of the configuration file as dictonary with the tag as
        key and the information as value
    '''
    file_path = file_in


    if file_path.endswith('.docx'):
        doc = Document(file_path)

        prop = doc.core_properties

        for replace_duo in config["ReplaceString"]:
            docx_replace(doc, replace_duo[0], replace_duo[1])
            prop.title = prop.title.replace(replace_duo[0], replace_duo[1])

        new_file_path = os.path.basename(file_path)
        doc.save((config["BetweenFolder"]) + new_file_path)

       # Open input document and new document
        with ZipFile(open((config["BetweenFolder"]) + new_file_path, "rb")) as zip_in:
            # copy document and replace content
            text_note = ''
            # check for logo
            status, note, warning = '', '', ''

            with ZipFile(file_out, "w", ZIP_DEFLATED) as zip_out:
                # copy document and replace content
                copy_and_replace(zip_in, zip_out, config)
                # check for logo
                status, note, warning = place_logo(zip_in, zip_out, config)
                # Add missing images
                add_missing_images(zip_in, zip_out)

        # Remove file from betweenFolder
        os.remove((config["BetweenFolder"]) + new_file_path)

        log.write(f"{file_in};{status};{note};{text_note};{warning}\n")

def add_missing_images(zip_in, zip_out):
    """
    Adds missing images from one Zip archive to another.

    Parameters
    ----------
    zip_in : ZipFile
        The input ZipFile object from which to extract the image filenames.

    zip_out : ZipFile
        The output ZipFile object where the missing images will be added.
    """
    # Get a list of image filenames from zip_in
    image_files = [item.filename for item in zip_in.infolist() if item.filename.startswith('word/media/')]

    # Check if the image already exists in zip_out, and if not, add it
    for image_file in image_files:
        if image_file not in zip_out.namelist():
            image_data = zip_in.read(image_file)
            zip_out.writestr(image_file, image_data)

def prepare_log(config):
    '''
    Creates and prepares logfile

    Parameters
    ----------
    config: dictionary
        content of the configuration file as dictonary with the tag as
        key and the information as value

    Returns
    -------
    log: TextIOWrapper
        wrapper for the logfile
    '''
    log_time = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    log_name = f"BH_Rebrand_{log_time}.csv"
    log_path = os.path.join(config["LogFolder"], log_name)

    # pylint: disable=R1732
    # Disables consider using with for resource allocation
    # Log needs to be closed in main function after processing all files
    log = open(log_path, "w+", encoding="utf-8")

    log.write("sep=;\n")
    log.write("Inputfile;Logo;Notes;LegacyText\n")

    print(f"Logging in file {log_name}")
    return log


def main():
    '''
    Starts processing of files and conversion to PDF
    '''
    if len(sys.argv) == 2:
        # Put elements of config file into a dictionary
        config = map_config(sys.argv[1])

        # Check PDF conversion
        pdf_conversion = False
        if "true" in config["PDF"].lower():
            pdf_conversion = True
            print("Starting COM Server for PDF conversion")
            # Start COM servers
            start_time = time()
            # Get a list of the filetypes that are going to be rebranded
            filetypes = get_filetypes_in_folder(config["InputFolder"])

            # Start the necessary COM Servers
            word, excel, power_point = start_com_servers(filetypes)

            # End timer and output time
            print(f"Starting server took {time() - start_time:.3f} seconds")


        # Check if output folder exists and create it if neccessary
        if not os.path.isdir(config["OutputFolder"]):
            os.mkdir(config["OutputFolder"])

        # Check if log folder exists and create it if neccessary
        if not os.path.isdir(config["LogFolder"]):
            os.mkdir(config["LogFolder"])

        # Check if input directory exists
        if os.path.isdir(config["InputFolder"]):
            log = prepare_log(config)

            # Loop over every file in the directory
            for file in os.listdir(config["InputFolder"]):
                # Start timer
                start_time = time()

                # Create input and output path and start file processing
                file_in = os.path.join(config["InputFolder"], file)
                file_out = os.path.join(config["OutputFolder"], file)

                # Check if current file is a path to a folder
                if os.path.isdir(file_in):
                    # Ignore folder and continue
                    continue

                # Get the current filetype
                config = get_filetype(file, config)

                # Check if filetype is supported
                if "filetype" not in config:
                    print(f"Filetype of {file} not supported")
                    log.write(f"{file_in};-;Filetype not supported\n")
                    continue

                # Format the OldLogoPath
                config["OldLogoPath_formatted"] = config["OldLogoPath"].format(
                    filetype = config["filetype"])
                config["LegacyBHLogoPath_formatted"] = config["LegacyBHLogoPath"].format(
                    filetype = config["filetype"])

                # Process the file
                print(f"Processing {file}")
                process_file(file_in, file_out, config, log)

                # End timer and output time
                print(f"Processing took {time() - start_time:.3f} seconds")

                # Convert to PDF
                if pdf_conversion:
                    # Start timer
                    start_converting_time = time()

                    # Convert file using COM server
                    print(f"Converting {file} to PDF")
                    config["PDFPath"] = os.path.join(config["OutputFolder"],
                                                     file[0:file.rfind(".")] + ".pdf")

                    # Convert the file to pdf
                    convert_to_pdf(file_out, config, word, excel, power_point)


                    # End timer and output time
                    print(f"Converting took {time() - start_converting_time:.3f} seconds")
            log.close()
        else:
            print("Specify an existing input folder containing the documents and an output folder")

        # Close COM Server
        if pdf_conversion:
            quit_com_servers(word, excel, power_point)

        print("All done")

    else:
        print("Specify an existing config file")


def map_config(configfile):
    """
    This function creates a dictionary from the information gathered form
    the configuration file.

    Parameters
    ----------
    configfile: string
        the path to the configuration file

    Returns
    -------
    cfg_dict: dictonary
        content of the configuration file as dictonary with the tag as
        key and the information as value

    """
    cfg_dict = {}
    cfg_dict["ReplaceString"] = []
    cfg_dict["case-unsensitive"] = []

    with open(configfile, encoding="utf-8") as conf:
        for line in conf:
            if not line.startswith("//"):
                if search(r'[\S\s]*? = [\S\s]*?', line) is not None:
                    (key, val)= line.split(' = ')
                    cfg_dict[key] = val.strip()
                elif search(r'[\S\s]*? -> [\S\s]*?', line) is not None:
                    line = str(line)
                    'line = bytes(line,  "utf-8")'
                    (old, new) = line.split(' -> ')
                    # pylint: disable=W1401
                    # Disable error from using \S\s in a binary string which is needed for regex
                    #if search(b'[\S\s]*? // case-unsensitive', new) is not None:
                    #    new = (new[0:new.find(b" // case-unsensitive")])
                    #    cfg_dict["case-unsensitive"].append(new)
                    cfg_dict["ReplaceString"].append([old.strip(), new.strip()])

    return cfg_dict


if __name__ == "__main__":
    main()